# Ensure proper module imports by fixing the Python path at the very start
import os
import sys

# Add the project root directory to Python path
current_file_path = os.path.abspath(__file__)
current_dir = os.path.dirname(current_file_path)
project_root = os.path.dirname(current_dir)

if current_dir not in sys.path:
    sys.path.insert(0, current_dir)
if project_root not in sys.path:
    sys.path.insert(0, project_root)

# Try to load environment variables from .env file
try:
    from dotenv import load_dotenv
    env_path = os.path.join(project_root, '.env')
    if os.path.exists(env_path):
        load_dotenv(env_path)
        print(f"✓ Loaded environment from {env_path}")
except ImportError:
    # python-dotenv not installed, will use defaults
    pass

# Standard library imports
import pathlib
import html
import re
import json
import datetime
import uuid
import time
import logging
import threading
import random
import base64
import mimetypes
from functools import wraps
import traceback
from io import BytesIO, StringIO
from urllib.parse import unquote
import subprocess
import platform

# Third-party imports
import requests
import flask
from flask import Flask, request, render_template, redirect, url_for, flash, session, jsonify, send_from_directory, send_file, make_response, abort
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from waitress import serve
    USE_WAITRESS = True
except ImportError:
    USE_WAITRESS = False

# Set up logging to a rotating file handler in /logs directory
log_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'logs')
os.makedirs(log_dir, exist_ok=True)

# Configure the root logger
from logging.handlers import RotatingFileHandler

# Create a formatter with timestamps
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')

# Set up file handler (rotating log, max 5MB, keep 5 backup files)
log_file_path = os.path.join(log_dir, f'anymize_{datetime.datetime.now().strftime("%Y%m%d")}.log')
file_handler = RotatingFileHandler(
    log_file_path,
    maxBytes=5*1024*1024,  # 5MB
    backupCount=5
)
file_handler.setFormatter(formatter)
file_handler.setLevel(logging.DEBUG)  # Capture ALL log levels

# Set up console handler
console_handler = logging.StreamHandler()
console_handler.setFormatter(formatter)
console_handler.setLevel(logging.INFO)  # More selective for console output

# Configure root logger
root_logger = logging.getLogger()
root_logger.setLevel(logging.DEBUG)  # Capture ALL log levels
for handler in root_logger.handlers[:]:  # Remove existing handlers
    root_logger.removeHandler(handler)
root_logger.addHandler(file_handler)
root_logger.addHandler(console_handler)

logging.info(f"Starting Anymize with logging to {log_file_path}")
logging.debug(f"Debug logging enabled - all messages will be logged to file")

# Initialize Flask app
app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB max upload size

# Initialize rate limiter
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://"
)

# Security: Generate secure secret key if not provided
if not os.environ.get('SECRET_KEY'):
    logging.warning("No SECRET_KEY found in environment. Generating a secure random key.")
    import secrets
    app.secret_key = secrets.token_hex(32)
else:
    app.secret_key = os.environ.get('SECRET_KEY')

# Session security configuration
app.config['SESSION_COOKIE_SECURE'] = True  # Only send cookies over HTTPS
app.config['SESSION_COOKIE_HTTPONLY'] = True  # Prevent JavaScript access to session cookie
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # CSRF protection
app.config['SESSION_COOKIE_NAME'] = 'anymize_session'  # Custom session cookie name

# Session timeout configuration - ONLY inactivity based
# Session persists indefinitely as long as user is active
# Expires after 4 hours of inactivity
from datetime import timedelta
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=365)  # Effectively no absolute limit
app.config['SESSION_REFRESH_EACH_REQUEST'] = True  # Refresh session expiry on activity
app.config['SESSION_INACTIVITY_TIMEOUT'] = timedelta(hours=4)  # Logout after 4 hours inactivity

# Configure Flask's built-in logging to use our logger
app.logger.handlers = []
for handler in logging.getLogger().handlers:
    app.logger.addHandler(handler)
app.logger.setLevel(logging.DEBUG)

# Import required modules and functions
try:
    # Create necessary directories function
    def ensure_directories():
        """Create necessary directories"""
        os.makedirs(os.path.join(app.root_path, 'uploads'), exist_ok=True)
        os.makedirs(os.path.join(app.root_path, 'static', 'temp'), exist_ok=True)
    
    logging.info("Using external OCR service for document processing")

    # Import API Blueprint
    try:
        # First try relative import (if run as module)
        from .api import api_bp
        logging.info("Imported API using relative import")
    except ImportError as e:
        logging.warning(f"Relative import failed: {e}, trying direct import")
        try:
            # Fall back to absolute import (if run as script)
            import sys
            sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            from anymize.api import api_bp
            logging.info("Imported API using absolute package import")
        except ImportError as e2:
            logging.warning(f"Package import failed: {e2}, trying direct import")
            try:
                # Last resort - try direct import
                import api
                api_bp = api.api_bp
                logging.info("Imported API using direct import")
            except ImportError as e3:
                logging.critical(f"All import methods failed! Application will not function correctly: {e3}")
                raise

except Exception as e:
    logging.critical(f"Error during initialization: {e}")
    raise

# Register the API blueprint
app.register_blueprint(api_bp, url_prefix='/api')

# Enable Flask error handling for ALL routes
@app.errorhandler(Exception)
def handle_exception(e):
    """Global exception handler for all routes"""
    error_id = str(uuid.uuid4())[:8]
    logging.error(f"Unhandled exception [{error_id}]: {str(e)}", exc_info=True)
    return jsonify({
        'error': f"Internal server error ({error_id})",
        'message': str(e),
        'status': 'error'
    }), 500

from config_shared import (
    get_job, 
    replace_prefixes_with_labels, 
    detect_language_and_get_prompt,
    OCR_WEBHOOK_URL,
    RAW_TEXT_WEBHOOK_URL,
    FURTHER_ANONYMIZATION_WEBHOOK_URL,
    N8N_WEBHOOK_URL,
    JOB_CREATE_ENDPOINT,
    JOB_UPDATE_ENDPOINT,
    HEADERS,
    TABLE_ID,
    NOCODB_BASE,
    JOB_GET_ENDPOINT,
    SYSTEM_PROMPTS,
    DOC_PROMPT_BEGIN,
    DOC_PROMPT_END,
    link_job_to_user,
    check_job_user_link
)

# Import authentication module
from auth import send_email_verification, verify_code, is_authenticated, get_current_user, logout as auth_logout_func
from functools import wraps

# Authentication decorator
def auth_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not is_authenticated():
            # Store the URL user wanted to access
            session['next_url'] = request.url
            # Only show error message if user was trying to access a protected page
            # and not coming directly to login or from logout
            if request.endpoint not in ['auth_login', 'auth_verify', 'auth_send_code', 'static']:
                flash('Bitte melden Sie sich an, um fortzufahren.', 'info')
            return redirect(url_for('auth_login'))
        
        # Check for session inactivity
        last_activity = session.get('last_activity')
        if last_activity:
            try:
                last_activity_time = datetime.datetime.fromisoformat(last_activity)
                inactivity_timeout = app.config.get('SESSION_INACTIVITY_TIMEOUT', timedelta(hours=4))
                
                # Check if session has been inactive too long
                if datetime.datetime.now() - last_activity_time > inactivity_timeout:
                    logging.info(f"Session timeout for user {session.get('user_email')} due to inactivity")
                    auth_logout_func()  # Clear session
                    flash('Ihre Sitzung ist aufgrund von Inaktivität abgelaufen. Bitte melden Sie sich erneut an.', 'warning')
                    session['next_url'] = request.url
                    return redirect(url_for('auth_login'))
                    
            except Exception as e:
                logging.error(f"Error checking session activity: {e}")
        
        # Update last activity timestamp
        session['last_activity'] = datetime.datetime.now().isoformat()
        
        return f(*args, **kwargs)
    return decorated_function

# Handle imports differently depending on how the script is run
try:
    # Try relative import first (when running as a module)
    from .api import trigger_n8n_webhook_async
except ImportError:
    try:
        # Try direct import (when running as script)
        from api import trigger_n8n_webhook_async
    except ImportError as e:
        logging.error(f"Failed to import trigger_n8n_webhook_async: {e}")
        # Define a fallback function if import fails
        def trigger_n8n_webhook_async(payload, record_id, job_id, webhook_url=None):
            """Fallback function when import fails"""
            logging.warning("Using fallback webhook trigger function")
            target_url = webhook_url or N8N_WEBHOOK_URL
            try:
                resp = requests.post(target_url, json=payload)
                logging.info(f"Direct webhook call: {resp.status_code}")
            except Exception as e:
                logging.error(f"DIRECT WEBHOOK ERROR: {str(e)}", exc_info=True)
            return True

# Import required packages for PDF and DOCX
try:
    # Dynamic import to handle runtime installation of reportlab
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import A4
except ImportError as e:
    logging.error(f"Reportlab import failed: {e}", exc_info=True)
    PDF_AVAILABLE = False
else:
    PDF_AVAILABLE = True

try:
    # Dynamic import to handle runtime installation of python-docx
    from docx import Document
    import docx.shared
except ImportError as e:
    logging.error(f"python-docx import failed: {e}", exc_info=True)
    DOCX_AVAILABLE = False
else:
    DOCX_AVAILABLE = True

# Track anonymization triggers per job
triggered_anon = set()
# Lock to synchronize anonymization triggers
anon_lock = threading.Lock()

# Serve Swagger UI at /swagger
@app.route('/swagger')
def swagger_ui():
    return '''
    <!DOCTYPE html>
    <html lang="en">
    <head>
      <meta charset="UTF-8">
      <title>Swagger UI</title>
      <link rel="stylesheet" type="text/css" href="https://unpkg.com/swagger-ui-dist/swagger-ui.css" />
    </head>
    <body>
      <div id="swagger-ui"></div>
      <script src="https://unpkg.com/swagger-ui-dist/swagger-ui-bundle.js"></script>
      <script>
        window.onload = function() {
          window.ui = SwaggerUIBundle({
            url: '/swagger.yaml',
            dom_id: '#swagger-ui',
          });
        };
      </script>
    </body>
    </html>
    '''

# Serve swagger.yaml statically
@app.route('/swagger.yaml')
def swagger_yaml():
    yaml_path = pathlib.Path(__file__).parent / 'swagger.yaml'
    return send_from_directory(yaml_path.parent, yaml_path.name, mimetype='text/yaml')

# Add swagger_url variable for templates
app.config['SWAGGER_URL'] = '/swagger'

# Helper function for generating secure 5-digit job IDs
def generate_unique_job_id():
    """Generates a 5-digit, numeric job ID."""
    return ''.join([str(random.randint(0, 9)) for _ in range(5)])

# Allowed file extensions and MIME types for security
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'gif', 'txt', 'doc', 'docx'}
ALLOWED_MIME_TYPES = {
    'application/pdf',
    'image/png', 
    'image/jpeg',
    'image/gif',
    'text/plain',
    'application/msword',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def validate_file_upload(file):
    """Validate uploaded file for security"""
    if not file or not file.filename:
        return False, "Keine Datei ausgewählt."
    
    # Check file extension
    if not allowed_file(file.filename):
        return False, f"Dateityp nicht erlaubt. Erlaubte Typen: {', '.join(ALLOWED_EXTENSIONS)}"
    
    # Check file size (read first bytes to get size without loading entire file)
    file.seek(0, 2)  # Seek to end
    file_size = file.tell()
    file.seek(0)  # Reset to beginning
    
    if file_size > MAX_FILE_SIZE:
        return False, f"Datei zu groß. Maximale Größe: {MAX_FILE_SIZE // (1024*1024)} MB"
    
    # Check MIME type
    file_content = file.read(1024)  # Read first 1KB
    file.seek(0)  # Reset position
    
    # Basic MIME type checking based on file signature
    mime_type = None
    if file_content.startswith(b'%PDF'):
        mime_type = 'application/pdf'
    elif file_content.startswith(b'\x89PNG'):
        mime_type = 'image/png'
    elif file_content.startswith(b'\xff\xd8\xff'):
        mime_type = 'image/jpeg'
    elif file_content.startswith(b'GIF8'):
        mime_type = 'image/gif'
    elif file_content.startswith(b'PK\x03\x04'):  # DOCX
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif all(c < 128 for c in file_content[:100] if c != 0):  # Basic text file check
        mime_type = 'text/plain'
    
    if mime_type and mime_type not in ALLOWED_MIME_TYPES:
        return False, "Dateityp nicht erlaubt basierend auf Dateiinhalt."
    
    return True, "OK"

# Authentication Routes
@app.route('/login')
def auth_login():
    """Show login page."""
    # If already authenticated, redirect to main page
    if is_authenticated():
        return redirect(url_for('index'))
    return render_template('auth_login.html')

@app.route('/auth/send-code', methods=['POST'])
@limiter.limit("5/hour")  # Strict limit for code sending
def auth_send_code():
    """Send verification code to email."""
    email = request.form.get('email', '').strip()
    
    if not email:
        flash('Bitte geben Sie eine E-Mail-Adresse ein.', 'error')
        return redirect(url_for('auth_login'))
    
    # Basic email validation
    if '@' not in email or '.' not in email:
        flash('Bitte geben Sie eine gültige E-Mail-Adresse ein.', 'error')
        return redirect(url_for('auth_login'))
    
    # Send verification code
    success = send_email_verification(email)
    
    if success:
        # Always redirect to verification page, even if database operations failed
        return redirect(url_for('auth_verify'))
    else:
        flash('Es gab ein Problem beim Senden der E-Mail. Bitte versuchen Sie es erneut.', 'error')
        return redirect(url_for('auth_login'))

@app.route('/auth/verify', methods=['GET', 'POST'])
@limiter.limit("100/hour")  # Increased limit for verification attempts
def auth_verify():
    """Verify the 6-digit code."""
    if request.method == 'GET':
        # Get email from session
        email = session.get('auth_email')
        if not email:
            flash('Bitte fordern Sie zuerst einen Verifizierungscode an.', 'error')
            return redirect(url_for('auth_login'))
        return render_template('auth_verify.html', email=email)
    
    # POST method - verify code
    email = session.get('auth_email')
    code = request.form.get('code', '').strip()
    
    if not email or not code:
        flash('Fehlende E-Mail oder Code.', 'error')
        return render_template('auth_verify.html', email=email)
    
    # Verify the code
    if verify_code(email, code):
        # Get the redirect URL from session or default to index
        next_url = session.pop('next_url', None)
        return redirect(next_url or url_for('index'))
    else:
        flash('Ungültiger Code. Bitte versuchen Sie es erneut.', 'error')
        return render_template('auth_verify.html', email=email)

@app.route('/auth/resend', methods=['POST'])
@limiter.limit("3/hour")  # Very strict limit for resend
def auth_resend_code():
    """Resend verification code."""
    email = request.form.get('email', '').strip()
    
    if not email:
        flash('E-Mail-Adresse fehlt.', 'error')
        return redirect(url_for('auth_login'))
    
    # Resend verification code
    if send_email_verification(email):
        flash('Ein neuer Code wurde an Ihre E-Mail gesendet.', 'success')
        return render_template('auth_verify.html', email=email)
    else:
        flash('Es ist ein Fehler aufgetreten. Bitte versuchen Sie es erneut.', 'error')
        return redirect(url_for('auth_login'))

@app.route('/logout')
def auth_logout():
    """Logout user."""
    auth_logout_func()
    flash('Sie wurden erfolgreich abgemeldet.', 'success')
    return redirect(url_for('auth_login'))

# Endpoints

def poll_for_output_text(job_id, max_retries=30, retry_delay=5):
    """Repeatedly query for full_prefix_text until it's filled or max_retries is reached"""
    def poll():
        for attempt in range(max_retries):
            job_data = get_job(job_id)
            if job_data and job_data.get("full_prefix_text"):
                logging.info(f"Full prefix text found for job {job_id}: {len(job_data['full_prefix_text'])} characters")
                break
            logging.info(f"Waiting for full_prefix_text for job {job_id} (attempt {attempt+1}/{max_retries})...")
            time.sleep(retry_delay)
    thread = threading.Thread(target=poll)
    thread.daemon = True
    thread.start()

@app.route('/upload', methods=['POST'])
@auth_required
@limiter.limit("10/minute")
def upload_file():
    """Handle file upload and redirection to result page"""
    # Validate file part
    if 'file' not in request.files or not request.files['file'].filename:
        flash('No file selected.', 'error')
        return redirect(url_for('index'))  # Changed from request.url to index route
    file = request.files['file']

    # Validate file upload
    is_valid, message = validate_file_upload(file)
    if not is_valid:
        flash(message, 'error')
        return redirect(url_for('index'))

    # Save upload
    filename = secure_filename(file.filename)
    upload_folder = os.path.join(app.root_path, 'uploads')
    os.makedirs(upload_folder, exist_ok=True)
    
    # Generate unique filename to prevent overwrites and path traversal
    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    random_suffix = str(uuid.uuid4())[:8]
    safe_filename = f"{timestamp}_{random_suffix}_{filename}"
    file_path = os.path.join(upload_folder, safe_filename)
    
    # Additional path validation
    if not os.path.abspath(file_path).startswith(os.path.abspath(upload_folder)):
        flash('Ungültiger Dateiname.', 'error')
        return redirect(url_for('index'))
    
    file.save(file_path)
    # Log with sanitized filename to prevent log injection
    safe_log_filename = filename.replace('\n', '\\n').replace('\r', '\\r')[:100]
    logging.info(f"File '{safe_log_filename}' saved as '{safe_filename}' at '{file_path}'")

    # Create job record in NoCodeDB with placeholder text
    job_id = str(uuid.uuid4())
    logging.info(f"Created job_id: {job_id}")
    job_payload = {
        'internal_ID': job_id,
        'input_text': 'Document submitted for OCR processing. Results will be available soon.',
        'status': 'processing'
    }
    try:
        resp = requests.post(JOB_CREATE_ENDPOINT, json=job_payload, headers=HEADERS)
        resp.raise_for_status()
        data = resp.json().get('data', resp.json())
        record_id = data.get('Id')
        session['job_id'] = record_id
        
        # Link job to user if authenticated
        user_id = session.get('user_id')
        if user_id:
            link_job_to_user(record_id, user_id)
            
    except Exception as e:
        logging.error(f"Failed to create job record: {e}", exc_info=True)
        flash('Error creating job. Please try again later.', 'error')
        return redirect(url_for('index'))

    # Prepare file copy for OCR processing
    temp_file_path = os.path.join(upload_folder, f"{job_id}_{filename}")
    if file_path != temp_file_path:
        import shutil
        shutil.copy2(file_path, temp_file_path)
        logging.info(f"Copied file to temp path {temp_file_path}")

    # Invoke OCR webhook
    file_handle = None
    try:
        file_handle = open(temp_file_path, 'rb')
        files = {
            'file': (filename, file_handle, mimetypes.guess_type(temp_file_path)[0] or 'application/octet-stream')
        }
        data = {'job_id': str(record_id)}
        ocr_resp = requests.post(OCR_WEBHOOK_URL, files=files, data=data, timeout=120)
        ocr_resp.raise_for_status()
        logging.info(f"OCR webhook called for job {record_id}")
    except Exception as e:
        logging.error(f"OCR webhook call failed: {e}", exc_info=True)
        flash('Error initiating OCR. Please try again.', 'error')
        return redirect(url_for('index'))
    finally:
        # Close file handle
        if file_handle:
            file_handle.close()
        
        # Clean up uploaded files after OCR processing
        try:
            # Delete the original uploaded file
            if os.path.exists(file_path):
                os.remove(file_path)
                logging.info(f"Deleted original file: {file_path}")
            
            # Delete the temporary file if it's different from the original
            if temp_file_path != file_path and os.path.exists(temp_file_path):
                os.remove(temp_file_path)
                logging.info(f"Deleted temp file: {temp_file_path}")
        except Exception as cleanup_error:
            logging.error(f"Failed to clean up files: {cleanup_error}")
    
    # Redirect using UUID param to hide primary ID
    return redirect(url_for('result', uuid=job_id))

@app.route('/upload_text', methods=['POST'])
@auth_required  
def upload_text():
    """Handle raw text input and send to webhook for processing"""
    # Get raw text from form
    raw_text = request.form.get('raw_text', '').strip()
    if not raw_text:
        flash('No text provided.', 'error')
        return redirect(url_for('index'))
    
    # Create job record in NoCodeDB
    job_id = str(uuid.uuid4())
    logging.info(f"Created job_id for text input: {job_id}")
    job_payload = {
        'internal_ID': job_id,
        'input_text': raw_text,
        'status': 'processing'
    }
    
    try:
        # Create job record
        resp = requests.post(JOB_CREATE_ENDPOINT, json=job_payload, headers=HEADERS)
        resp.raise_for_status()
        data = resp.json().get('data', resp.json())
        record_id = data.get('Id')
        session['job_id'] = record_id
        logging.info(f"Created job record with ID: {record_id}")
        
        # Link job to user if authenticated
        user_id = session.get('user_id')
        if user_id:
            link_job_to_user(record_id, user_id)
            
    except Exception as e:
        logging.error(f"Failed to create job record: {e}", exc_info=True)
        flash('Error creating job. Please try again later.', 'error')
        return redirect(url_for('index'))
    
    # Send text to the webhook for processing
    try:
        webhook_url = RAW_TEXT_WEBHOOK_URL
        webhook_payload = {
            'text': raw_text,
            'job_id': str(record_id)
        }
        
        logging.info(f"Sending raw text to webhook: {webhook_url}")
        webhook_resp = requests.post(webhook_url, json=webhook_payload, timeout=120)
        webhook_resp.raise_for_status()
        logging.info(f"Text processing webhook called successfully for job {record_id}")
    except Exception as e:
        logging.error(f"Text processing webhook call failed: {e}", exc_info=True)
        flash('Error processing text. Please try again.', 'error')
        return redirect(url_for('index'))
    
    # Redirect to result page
    return redirect(url_for('result', uuid=job_id))

@app.route('/', methods=['GET', 'POST'])
@auth_required
def index():
    """Main route for file upload and processing"""
    if request.method == 'POST':
        if 'file' not in request.files or request.files['file'].filename == '':
            flash('No file selected.', 'error')
            return redirect(request.url)
        file = request.files['file']

        # Validate file upload
        is_valid, message = validate_file_upload(file)
        if not is_valid:
            flash(message, 'error')
            return redirect(url_for('index'))

        # Save file temporarily
        upload_folder = 'uploads'
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder)
        filename = secure_filename(file.filename)
        # Generate unique filename to prevent overwrites
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        random_suffix = str(uuid.uuid4())[:8]
        safe_filename = f"{timestamp}_{random_suffix}_{filename}"
        file_path = os.path.join(upload_folder, safe_filename)
        
        # Additional path validation
        if not os.path.abspath(file_path).startswith(os.path.abspath(upload_folder)):
            flash('Ungültiger Dateiname.', 'error')
            return redirect(url_for('index'))
        
        file.save(file_path)
        logging.info(f"File '{filename}' saved at '{file_path}'")

        try:
            # Create job in NoCodeDB - initially without text (will be filled by OCR service)
            job_id = generate_unique_job_id()
            job_payload = {
                'internal_ID': job_id,  # Correct field name per NocoDB v2 schema
                'output_text': '',  # Initially empty, will be filled by LLM
                'status': 'processing'
            }

            # Create the job entry in NocoDB
            resp = requests.post(JOB_CREATE_ENDPOINT, json=job_payload, headers=HEADERS)
            if resp.status_code == 200:
                resp_json = resp.json()
                record_id = resp_json.get('Id')
                logging.info(f"NoCodeDB Job Create Response: {resp.status_code} {resp.text}")
                logging.info(f"Job created, ID: {record_id}")

                # Store job ID in session
                session['job_id'] = record_id
                
                # Link job to user if authenticated
                user_id = session.get('user_id')
                if user_id:
                    link_job_to_user(record_id, user_id)

                # Create temporary file path to ensure it exists during processing
                temp_file_path = os.path.join(upload_folder, f"{job_id}_{filename}")
                if file_path != temp_file_path:
                    # Copy the file to ensure it exists during async processing
                    import shutil
                    shutil.copy2(file_path, temp_file_path)
                    logging.info(f"Created temporary file for processing: {temp_file_path}")
                
                # DO NOT send to OCR webhook here - only do it in the upload_file route
                # THIS LINE WAS CAUSING DUPLICATE OCR CALLS
                logging.info(f"Skipping OCR call in index route - upload_file will handle it")

                # Update the job with initial input text as a placeholder
                update_resp = requests.post(
                    f"{JOB_UPDATE_ENDPOINT}/{record_id}",
                    json={'input_text': "Document sent for OCR processing. Results will be available soon."},
                    headers=HEADERS
                )
                
                # Log success message
                logging.info(f"Document successfully sent to OCR service for job {record_id}")
                
                # Wait for OCR processing to complete or timeout
                max_wait_time = 120  # 120 seconds timeout
                ocr_successful = False
                
                # Show a loading message on the upload page
                flash('Processing document, please wait...', 'info')
                
                # Wait for OCR results with timeout
                start_time = time.time()
                while time.time() - start_time < max_wait_time:
                    # Check if OCR has completed by looking for input_text in the job
                    job_data = get_job(record_id, log_request=False)
                    
                    if job_data and job_data.get('input_text') and len(job_data.get('input_text', '')) > 30:
                        # OCR has completed successfully
                        logging.info(f"OCR completed for job {record_id}, redirecting to result page")
                        ocr_successful = True
                        
                        # Now we can trigger the anonymization workflow using the extracted text
                        extracted_text = job_data.get('input_text')
                        n8n_payload = {
                            'id': record_id,
                            'internal_ID': job_id,
                            'text': extracted_text,
                            'action': 'process',
                            'char_count': len(extracted_text)
                        }
                                                # Use the anonymization webhook from config
                        # Using the imported webhook URL
                        EXACT_WEBHOOK_URL = N8N_WEBHOOK_URL
                        
                        logging.info(f"ANONYMIZATION: Using exact webhook URL: {EXACT_WEBHOOK_URL}")
                        logging.info(f"ANONYMIZATION: Text length: {len(extracted_text)}, job: {record_id}")
                        
                        # Simple payload with only the necessary fields
                        simple_payload = {
                            'id': record_id,
                            'text': extracted_text
                        }
                        
                        def call_anon_webhook_directly(url, payload):
                            """Attempt to call webhook with direct requests"""
                            try:
                                import requests
                                logging.info(f"DIRECT WEBHOOK CALL to {url} with payload keys: {list(payload.keys())}")
                                
                                # Make the most basic request possible
                                headers = {'Content-Type': 'application/json', 'Accept': 'application/json'}
                                resp = requests.post(url, json=payload, headers=headers, timeout=10)
                                
                                logging.info(f"DIRECT WEBHOOK STATUS: {resp.status_code}")
                                return resp.status_code == 200
                            except Exception as e:
                                logging.error(f"DIRECT WEBHOOK ERROR: {str(e)}", exc_info=True)
                                return False
                        
                        # First try with simple payload
                        if call_anon_webhook_directly(EXACT_WEBHOOK_URL, simple_payload):
                            logging.info("ANONYMIZATION: Successfully triggered with simple payload")
                        else:
                            # Try with full payload as backup
                            full_payload = {
                                'id': record_id,
                                'internal_ID': job_id,
                                'text': extracted_text,
                                'action': 'process',
                                'char_count': len(extracted_text)
                            }
                            if call_anon_webhook_directly(EXACT_WEBHOOK_URL, full_payload):
                                logging.info("ANONYMIZATION: Successfully triggered with full payload")
                            else:
                                logging.error("ANONYMIZATION: Failed with both payload types")
                        
                        # Log that OCR is complete, but let the dedicated polling function handle the anonymization
                        logging.info(f"OCR processing complete for job {record_id} with {len(extracted_text)} chars of text")
                        logging.info(f"Sample text: {extracted_text[:100] if extracted_text else 'None'}")
                            
                        # Redirect to the result page since OCR is complete
                        flash('Document processed successfully!', 'success')
                        return redirect(url_for('result', uuid=job_id))
                    
                    # Sleep briefly before checking again
                    time.sleep(0.5)
                
                if not ocr_successful:
                    # OCR processing timeout - show error on the upload page
                    logging.error(f"OCR processing timeout for job {record_id}")
                    flash('Document processing is taking longer than expected. Please check results page later or try again.', 'warning')
                    return redirect(url_for('index'))
            else:
                logging.error(f"NoCodeDB error: {resp.status_code} {resp.text}")
                flash('Error creating job. Please try again later.', 'error')
                return redirect(url_for('index'))

        except Exception as e:
            logging.error(f"Error processing document: {str(e)}", exc_info=True)
            flash(f'Error processing: {str(e)}', 'error')
            return redirect(request.url)
    else:
        # GET request - show the upload form
        return render_template('index.html', current_user=get_current_user())

@app.route('/result/<uuid>')
@auth_required
def result(uuid):
    """Result page displaying OCR results"""
    # Clear any existing flash messages to prevent old messages from appearing
    session.pop('_flashes', None)
    
    # Get current user for display
    current_user = get_current_user()
    current_user_id = session.get('user_id')
    
    # Use the UUID to look up the job
    initial = get_job(uuid)
    if not initial:
        flash('Job not found', 'danger')
        return redirect(url_for('index'))
    
    job_id = initial['Id']  # Get the numeric ID
    
    # Security check: Verify that this job belongs to the current user
    if current_user_id:
        if not check_job_user_link(job_id, current_user_id):
            flash('Sie haben keine Berechtigung, diesen Job anzuzeigen.', 'danger')
            return redirect(url_for('index'))
    
    try:
        # Extract text from job data
        job_data = get_job(job_id)
        if not job_data:
            flash('Job not found', 'danger')
            return redirect(url_for('index'))
            
        # Get the uuid from job_data if not provided as parameter
        if not uuid:
            uuid = job_data.get('internal_ID')
            
        input_text = job_data.get('input_text') or ''
        # Use output_text (final anonymized) for display
        raw_text = job_data.get('output_text') or ''
        # Determine language (use provided or detect)
        lang_code = job_data.get('language') if job_data.get('language') in ['de','en','es','it','fr'] else detect_language_and_get_prompt(input_text)[1]
        output_text = replace_prefixes_with_labels(raw_text, lang_code)

        # Check if processing is complete
        processing = (raw_text == '')

        # Pass a timestamp to force template reload and avoid caching issues
        current_timestamp = int(time.time())

        return render_template(
            'enhanced_result.html',
            job_id=job_id,
            record_id=0,
            input_text=input_text,
            output_text=output_text,
            processing=processing,
            current_timestamp=current_timestamp,
            uuid=uuid,
            further_anonymization_webhook_url=FURTHER_ANONYMIZATION_WEBHOOK_URL,
            current_user=current_user
        )

    except Exception as e:
        logging.error(f"Error rendering result page: {str(e)}")
        flash(f'Error retrieving job data: {str(e)}', 'danger')
        return redirect(url_for('index'))

@app.route('/check_status', methods=['GET'])
def check_status():
    """API endpoint for checking processing status via session"""
    job_id = session.get('job_id')
    if not job_id:
        return jsonify({"status": "error", "message": "No job found in session"}), 404

    job_data = get_job(job_id)
    if not job_data:
        return jsonify({"status": "error", "message": "Job not found"}), 404

    output_text = job_data.get("output_text", "")
    if output_text:
        # Get language from job data or fall back to detection
        lang_code = job_data.get('language', '')
        
        # If language is not provided by the OCR service or is invalid, fall back to detection
        if not lang_code or lang_code not in ['de', 'en', 'es', 'it', 'fr']:
            logging.warning(f"Language not provided by OCR service or invalid: {lang_code}, falling back to detection")
            prompt, lang_code = detect_language_and_get_prompt(job_data.get("input_text", ""))
        else:
            logging.info(f"Using language from OCR service: {lang_code}")
        
        output_text_labeled = replace_prefixes_with_labels(output_text, lang_code)
        return jsonify({
            "status": "completed",
            "output_text": output_text_labeled
        })
    else:
        return jsonify({
            "status": "processing",
            "output_text": ""
        })

@app.route('/result_ajax', methods=['GET'])
@limiter.limit("300/minute")  # Allow more requests for polling
def result_ajax():
    """AJAX endpoint for fetching job results."""
    try:
        # Use numeric job_id for DB fetch
        job_id = request.args.get('job_id')
        if not job_id:
            return jsonify({"error": "No job_id provided"}), 400
        job_data = get_job(job_id)
        if not job_data:
            return jsonify({"error": "Job not found", "job_id": job_id}), 404
        # Use numeric record ID for internal triggers
        record_id = job_data.get('Id')
        # Extract texts
        input_text = job_data.get('input_text') or ''
        # Use output_text (final anonymized) for display
        raw_text = job_data.get('output_text') or ''
        lang_code = job_data.get('language') if job_data.get('language') in ['de','en','es','it','fr'] else detect_language_and_get_prompt(input_text)[1]
        output_text = replace_prefixes_with_labels(raw_text, lang_code)

        logging.info(f"[AJAX] Polling for job {record_id}")

        # Trigger anonymization when OCR text ready (once)
        placeholder = "Document submitted for OCR processing. Results will be available soon."
        if input_text and input_text != placeholder:
            with anon_lock:
                if record_id not in triggered_anon:
                    # Avoid duplicate anonymization calls
                    triggered_anon.add(record_id)
                    logging.info(f"[ANON_TRIGGER] OCR text ready; triggering anonymization for record {record_id}")
                    try:
                        resp = requests.post(N8N_WEBHOOK_URL, json={'id': record_id, 'text': input_text}, timeout=10)
                        logging.info(f"[ANON_TRIGGER] Response {resp.status_code}")
                    except Exception as e:
                        logging.error(f"[ANON_TRIGGER] Error: {e}", exc_info=True)
                else:
                    logging.debug(f"[ANON_TRIGGER] Already triggered for record {record_id}")
        else:
            logging.debug(f"[ANON_TRIGGER] OCR text not ready for job {job_id}")

        # Return JSON response with labeled anonymized text
        return jsonify({
            'job_id': record_id,
            'input_text': input_text,
            'output_text': output_text,
            'status': 'complete' if raw_text else 'processing'
        })

    except Exception as e:
        logging.error(f"[AJAX] Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/download/pdf/<job_id>', methods=['GET'])
@auth_required
def download_pdf(job_id):
    logging.info(f"[download_pdf] invoked for job {job_id}")
    # Dynamic import to handle runtime installation of reportlab
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.pagesizes import A4
    except ImportError as e:
        logging.error(f"Reportlab import failed: {e}", exc_info=True)
        flash('PDF generation not available on this server', 'error')
        # Get UUID for redirect
        job_data = get_job(job_id)
        uuid_val = job_data.get('internal_ID') if job_data else job_id
        return redirect(url_for('result', uuid=uuid_val))

    try:
        # Get job data
        job_data = get_job(job_id)
        if not job_data:
            flash('Job not found', 'error')
            return redirect(url_for('index'))
        
        # Security check: Verify user has permission to access this job
        current_user_id = session.get('user_id')
        if current_user_id:
            numeric_job_id = job_data.get('Id')
            if numeric_job_id and not check_job_user_link(numeric_job_id, current_user_id):
                flash('Sie haben keine Berechtigung, diesen Job herunterzuladen.', 'danger')
                return redirect(url_for('index'))

        output_text = job_data.get('output_text', '')
        if not output_text:
            # Fall back to full_prefix_text
            output_text = job_data.get('full_prefix_text', '')
        input_text = job_data.get('input_text', '')
        logging.info(f"[download_pdf] output_text length: {len(output_text)}, input_text length: {len(input_text)}")
        if not output_text:
            flash('No output text available for this job', 'error')
            uuid_val = job_data.get('internal_ID', job_id)
            return redirect(url_for('result', uuid=uuid_val))

        # Get language from job data or fall back to detection
        lang_code = job_data.get('language', '')
        
        # If language is not provided by the OCR service or is invalid, fall back to detection
        if not lang_code or lang_code not in ['de', 'en', 'es', 'it', 'fr']:
            logging.warning(f"Language not provided by OCR service or invalid: {lang_code}, falling back to detection")
            prompt, lang_code = detect_language_and_get_prompt(input_text)
        else:
            # Get the prompt for the detected language
            prompt = SYSTEM_PROMPTS.get(lang_code, SYSTEM_PROMPTS['en'])
            logging.info(f"Using language from OCR service: {lang_code}")

        # Get begin and end prompts for the document
        begin_prompt = DOC_PROMPT_BEGIN.get(lang_code, DOC_PROMPT_BEGIN['en'])
        end_prompt = DOC_PROMPT_END.get(lang_code, DOC_PROMPT_END['en'])

        # PDF generation
        # Removed static PDF_AVAILABLE check; dynamic import ensures availability

        # Generate PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Create a red style for the prompts
        red_style = styles['Normal'].clone('RedText')
        red_style.textColor = 'red'

        # Add begin prompt in red
        p_begin = Paragraph(begin_prompt.replace('\n', '<br/>'), red_style)
        story.append(p_begin)

        # Add separator line
        story.append(Paragraph("<br/>", styles['Normal']))
        story.append(Paragraph("<hr width='100%'/><br/>", styles['Normal']))
        story.append(Paragraph("<b>*** START DOCUMENT TEXT ***</b><br/>", red_style))
        story.append(Paragraph("<br/>", styles['Normal']))

        # Add anonymized text
        p = Paragraph(output_text.replace('\n', '<br/>'), styles['Normal'])
        story.append(p)

        # Add spacing and end separator
        story.append(Paragraph("<br/>", styles['Normal']))
        story.append(Paragraph("<b>*** END DOCUMENT TEXT ***</b><br/>", red_style))
        story.append(Paragraph("<hr width='100%'/><br/>", styles['Normal']))

        # Add end prompt in red
        p_end = Paragraph(end_prompt.replace('\n', '<br/>'), red_style)
        story.append(p_end)

        # Build the PDF into memory buffer
        doc.build(story)
        buffer.seek(0)
        logging.info(f"[download_pdf] story elements count: {len(story)}")
        pdf_data = buffer.getvalue()
        logging.info(f"[download_pdf] pdf header bytes: {repr(pdf_data[:8])}, size: {len(pdf_data)}")
        response = make_response(pdf_data)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename=anonymized_{job_id}.pdf'
        response.headers['Content-Length'] = str(len(pdf_data))
        return response

    except Exception as e:
        logging.error(f"Error generating PDF: {str(e)}", exc_info=True)
        flash(f'Error generating PDF: {str(e)}', 'error')
        # Get UUID for redirect
        job_data = get_job(job_id)
        uuid_val = job_data.get('internal_ID') if job_data else job_id
        return redirect(url_for('result', uuid=uuid_val))

@app.route('/download/docx/<job_id>', methods=['GET'])
@auth_required
def download_docx(job_id):
    logging.info(f"[download_docx] invoked for job {job_id}")
    # Dynamic import to handle runtime installation of python-docx
    try:
        from docx import Document
        import docx.shared
    except ImportError as e:
        logging.error(f"python-docx import failed: {e}", exc_info=True)
        flash('DOCX generation not available on this server', 'error')
        # Get UUID for redirect
        job_data = get_job(job_id)
        uuid_val = job_data.get('internal_ID') if job_data else job_id
        return redirect(url_for('result', uuid=uuid_val))
    
    """Generate and download anonymized text as DOCX file"""
    try:
        # Get job data
        job_data = get_job(job_id)
        if not job_data:
            flash('Job not found', 'error')
            return redirect(url_for('index'))
        
        # Security check: Verify user has permission to access this job
        current_user_id = session.get('user_id')
        if current_user_id:
            numeric_job_id = job_data.get('Id')
            if numeric_job_id and not check_job_user_link(numeric_job_id, current_user_id):
                flash('Sie haben keine Berechtigung, diesen Job herunterzuladen.', 'danger')
                return redirect(url_for('index'))

        output_text = job_data.get('output_text', '')
        if not output_text:
            # Fall back to full_prefix_text
            output_text = job_data.get('full_prefix_text', '')
        input_text = job_data.get('input_text', '')
        logging.info(f"[download_docx] output_text length: {len(output_text)}, input_text length: {len(input_text)}")
        if not output_text:
            flash('No output text available for this job', 'error')
            uuid_val = job_data.get('internal_ID', job_id)
            return redirect(url_for('result', uuid=uuid_val))

        # Get language from job data or fall back to detection
        lang_code = job_data.get('language', '')
        
        # If language is not provided by the OCR service or is invalid, fall back to detection
        if not lang_code or lang_code not in ['de', 'en', 'es', 'it', 'fr']:
            logging.warning(f"Language not provided by OCR service or invalid: {lang_code}, falling back to detection")
            prompt, lang_code = detect_language_and_get_prompt(input_text)
        else:
            # Get the prompt for the detected language
            prompt = SYSTEM_PROMPTS.get(lang_code, SYSTEM_PROMPTS['en'])
            logging.info(f"Using language from OCR service: {lang_code}")

        # Get begin and end prompts for the document
        begin_prompt = DOC_PROMPT_BEGIN.get(lang_code, DOC_PROMPT_BEGIN['en'])
        end_prompt = DOC_PROMPT_END.get(lang_code, DOC_PROMPT_END['en'])

        # DOCX generation
        # Removed static DOCX_AVAILABLE check; dynamic import ensures availability

        # Generate DOCX in memory
        buffer = BytesIO()
        document = Document()

        # Add begin prompt in red
        # Split begin prompt into paragraphs
        begin_paragraphs = begin_prompt.split('\n')
        for paragraph in begin_paragraphs:
            if paragraph.strip():
                p = document.add_paragraph()
                run = p.add_run(paragraph)
                run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # Red color

        # Add separator
        separator = document.add_paragraph()
        separator_run = separator.add_run("------------------------------------------")
        document.add_paragraph()
        start_marker = document.add_paragraph()
        start_run = start_marker.add_run("*** START DOCUMENT TEXT ***")
        start_run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # Red color
        start_run.bold = True
        document.add_paragraph()

        # Split text into paragraphs and add to document
        paragraphs = output_text.split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                document.add_paragraph(paragraph)

        # Add end separator
        end_marker = document.add_paragraph()
        end_run = end_marker.add_run("*** END DOCUMENT TEXT ***")
        end_run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # Red color
        end_run.bold = True
        document.add_paragraph()
        separator2 = document.add_paragraph()
        separator2_run = separator2.add_run("------------------------------------------")

        # Add end prompt in red
        # Split end prompt into paragraphs
        end_paragraphs = end_prompt.split('\n')
        for paragraph in end_paragraphs:
            if paragraph.strip():
                p = document.add_paragraph()
                run = p.add_run(paragraph)
                run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # Red color

        # Save DOCX into memory buffer
        document.save(buffer)
        buffer.seek(0)
        logging.info(f"[download_docx] document paragraphs count: {len(document.paragraphs)}")
        docx_data = buffer.getvalue()
        logging.info(f"[download_docx] docx header bytes: {repr(docx_data[:4])}, size: {len(docx_data)}")
        response = make_response(docx_data)
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename=anonymized_{job_id}.docx'
        response.headers['Content-Length'] = str(len(docx_data))
        return response

    except Exception as e:
        logging.error(f"Error generating DOCX: {str(e)}", exc_info=True)
        flash(f'Error generating DOCX: {str(e)}', 'error')
        # Get UUID for redirect
        job_data = get_job(job_id)
        uuid_val = job_data.get('internal_ID') if job_data else job_id
        return redirect(url_for('result', uuid=uuid_val))

@app.route('/download/text/<job_id>', methods=['GET'])
@auth_required
def download_text(job_id):
    """Download anonymized text as plain text file with prompts"""
    try:
        # Get job data
        job_data = get_job(job_id)
        if not job_data:
            flash('Job not found', 'error')
            return redirect(url_for('index'))
        
        # Security check: Verify user has permission to access this job
        current_user_id = session.get('user_id')
        if current_user_id:
            numeric_job_id = job_data.get('Id')
            if numeric_job_id and not check_job_user_link(numeric_job_id, current_user_id):
                flash('Sie haben keine Berechtigung, diesen Job herunterzuladen.', 'danger')
                return redirect(url_for('index'))

        output_text = job_data.get('output_text', '')
        input_text = job_data.get('input_text', '')
        if not output_text:
            # Fall back to full_prefix_text
            output_text = job_data.get('full_prefix_text', '')
        logging.info(f"[download_text] output_text length: {len(output_text)}, input_text length: {len(input_text)}")
        if not output_text:
            flash('No output text available for this job', 'error')
            uuid_val = job_data.get('internal_ID', job_id)
            return redirect(url_for('result', uuid=uuid_val))

        # Get language from job data or fall back to detection
        lang_code = job_data.get('language', '')
        
        # If language is not provided by the OCR service or is invalid, fall back to detection
        if not lang_code or lang_code not in ['de', 'en', 'es', 'it', 'fr']:
            logging.warning(f"Language not provided by OCR service or invalid: {lang_code}, falling back to detection")
            prompt, lang_code = detect_language_and_get_prompt(input_text)
        else:
            # Get the prompt for the detected language
            prompt = SYSTEM_PROMPTS.get(lang_code, SYSTEM_PROMPTS['en'])
            logging.info(f"Using language from OCR service: {lang_code}")

        # Get begin and end prompts for the document
        begin_prompt = DOC_PROMPT_BEGIN.get(lang_code, DOC_PROMPT_BEGIN['en'])
        end_prompt = DOC_PROMPT_END.get(lang_code, DOC_PROMPT_END['en'])

        # Combine text with prompts
        # Note: Plain text can't have colors, but we'll add a note to indicate the prompts should be red
        full_text = f"[RED TEXT START]\n{begin_prompt}\n[RED TEXT END]\n\n" + \
                  f"------------------------------------------\n" + \
                  f"[RED TEXT]*** START DOCUMENT TEXT ***[/RED TEXT]\n\n" + \
                  f"{output_text}\n\n" + \
                  f"[RED TEXT]*** END DOCUMENT TEXT ***[/RED TEXT]\n" + \
                  f"------------------------------------------\n\n" + \
                  f"[RED TEXT START]\n{end_prompt}\n[RED TEXT END]"

        # Create a response with the text content
        response = make_response(full_text)
        response.headers["Content-Disposition"] = f"attachment; filename=anonymized_{job_id}.txt"
        response.headers["Content-Type"] = "text/plain"

        return response

    except Exception as e:
        logging.error(f"Error generating text file: {str(e)}", exc_info=True)
        flash(f'Error generating text file: {str(e)}', 'error')
        # Get UUID for redirect
        job_data = get_job(job_id)
        uuid_val = job_data.get('internal_ID') if job_data else job_id
        return redirect(url_for('result', uuid=uuid_val))

@app.route('/download/original/<job_id>', methods=['GET'])
@auth_required
def download_original(job_id):
    """Download original extracted text as plain text file"""
    try:
        # Get job data
        job_data = get_job(job_id)
        if not job_data:
            flash('Job not found', 'error')
            return redirect(url_for('index'))
        
        # Security check: Verify user has permission to access this job
        current_user_id = session.get('user_id')
        if current_user_id:
            # Extract numeric job ID from job_data
            numeric_job_id = job_data.get('Id')
            if numeric_job_id and not check_job_user_link(numeric_job_id, current_user_id):
                flash('Sie haben keine Berechtigung, diesen Job herunterzuladen.', 'danger')
                return redirect(url_for('index'))

        input_text = job_data.get('input_text', '')
        if not input_text:
            flash('No input text available for this job', 'error')
            return redirect(url_for('result', uuid=job_id))

        # Create a response with the text content
        response = make_response(input_text)
        response.headers["Content-Disposition"] = f"attachment; filename=original_{job_id}.txt"
        response.headers["Content-Type"] = "text/plain"

        return response

    except Exception as e:
        logging.error(f"Error generating text file: {str(e)}", exc_info=True)
        flash(f'Error generating text file: {str(e)}', 'error')
        return redirect(url_for('result', uuid=job_id))

@app.route('/api/job/<job_id>/status', methods=['GET'])
def get_job_status(job_id):
    """Get the status of a job including anonymized text if available"""
    try:
        job_data = get_job(job_id)
        if not job_data:
            return jsonify({'error': 'Job not found'}), 404
        
        # Use output_text (final anonymized) for display
        raw_text = job_data.get('output_text') or ''
        lang_code = job_data.get('language') if job_data.get('language') in ['de','en','es','it','fr'] else detect_language_and_get_prompt(job_data.get('input_text', ''))[1]
        output_text = replace_prefixes_with_labels(raw_text, lang_code) if raw_text else None
        
        response = {
            'status': 'completed' if raw_text else 'processing',
            'output_text': output_text,
            'error': job_data.get('error', None)
        }
        
        return jsonify(response)
    except Exception as e:
        logging.error(f"Error getting job status: {str(e)}", exc_info=True)
        return jsonify({'error': 'Internal server error'}), 500

# Create a thread-safe lock for OCR webhook calls
import threading
ocr_lock = threading.Lock()

# Create a tracker for processing jobs
processed_jobs = set()

def extract_text_from_file(file_path, job_id):
    """Send document to OCR service with thread-safe locking to prevent duplicate calls.
    This function is the SINGLE point for sending documents to the OCR service.
    """
    # Use job_id as the unique identifier for this job
    # We use a global lock to ensure thread safety
    with ocr_lock:
        # Check if we've already processed this job
        if job_id in processed_jobs:
            logging.warning(f"🚫 BLOCKED DUPLICATE OCR: Job {job_id} already being processed")
            return "Document already submitted for OCR processing."
            
        # Mark this job as being processed BEFORE we do anything else
        # This prevents race conditions where multiple threads try to process the same job
        processed_jobs.add(job_id)
        logging.info(f"🔒 LOCKED OCR PROCESSING: Job {job_id} is now locked for processing")
    
    # Now we can safely process this job outside the lock
    try:
        # Verify the file exists
        if not os.path.exists(file_path):
            logging.error(f"❌ FILE NOT FOUND: {file_path}")
            processed_jobs.remove(job_id)  # Remove from processed set
            return "Error: File not found"
            
        # Get file info for logging
        try:
            file_size = os.path.getsize(file_path)
            filename = os.path.basename(file_path)
        except Exception as e:
            logging.error(f"⚠️ Error getting file info: {e}")
            file_size = "unknown"
            filename = os.path.basename(file_path) if file_path else "unknown"
        
        # Log the OCR request details
        logging.info(f"📝 OCR REQUEST DETAILS: job={job_id}, file={filename}, size={file_size}")
        
        # Get webhook URL directly from config to avoid any import issues
        from config_shared import OCR_WEBHOOK_URL
        webhook_url = OCR_WEBHOOK_URL
        
        # Prepare file with proper MIME type
        with open(file_path, 'rb') as file_content:
            # Detect MIME type
            mime_type = mimetypes.guess_type(filename)[0] or 'application/octet-stream'
            
            # Create files dict with correct content type
            files = {
                'file': (filename, file_content, mime_type)
            }
            
            # CRITICAL: Parameter MUST be named 'job_id' (not record_id)
            data = {
                'job_id': str(job_id)  # String format for consistency
            }
            
            # Log exactly what we're sending
            logging.info(f"📣 SENDING TO OCR WEBHOOK: url={webhook_url}, job_id={job_id}")
            
            # Make the HTTP request with appropriate timeouts
            ocr_response = requests.post(
                webhook_url,
                files=files,
                data=data,
                timeout=120  # 120 second timeout
            )
            
            # Process the response
            if ocr_response.status_code == 200:
                logging.info(f"✅ OCR WEBHOOK SUCCESS: job={job_id}, response={ocr_response.text}")
                # Clean up the file after successful OCR processing
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        logging.info(f"✅ CLEANUP: Deleted file after OCR: {file_path}")
                except Exception as cleanup_error:
                    logging.error(f"⚠️ CLEANUP ERROR: Failed to delete file {file_path}: {cleanup_error}")
                return "Document submitted for OCR processing. Results will be available soon."
            else:
                # Log the error
                logging.error(f"❌ OCR WEBHOOK FAILED: job={job_id}, status={ocr_response.status_code}, response={ocr_response.text}")
                return "Error during OCR processing. Please try again."
                
    except Exception as e:
        # Log any exceptions
        logging.error(f"❌ OCR PROCESSING ERROR: job={job_id}, error={str(e)}", exc_info=True)
        return "Error during OCR processing. Please try again."
        
    # NOTE: We intentionally don't remove job_id from processed_jobs
    # This ensures the job is never processed twice, even if there's an error

# This application now uses an external OCR service rather than local OCR methods
if __name__ == '__main__':
    # Ensure all necessary directories exist before starting the app
    ensure_directories()
    port = int(os.environ.get("PORT", 8000))  # Standardport für Ausführung hinter Apache
    print("\n\n--- ROUTES REGISTERED ---")
    for rule in app.url_map.iter_rules():
        print(rule)
    print("------------------------\n\n")
    # Startup info with public domain and port
    print(f"API info: GET http://0.0.0.0:{port}/api | POST http://0.0.0.0:{port}/api/anonymize")
    if USE_WAITRESS:
        serve(app, host="0.0.0.0", port=port)
    else:
        app.run(host="0.0.0.0", port=port, debug=True)
